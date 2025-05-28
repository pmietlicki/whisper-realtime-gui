import sys
import sounddevice as sd
import numpy as np
import whisper
import queue
import threading
import time
import torch
import traceback
import math
import re
import os
from datetime import datetime
from PySide6.QtCore import Qt, QTimer, Signal, QThread, Slot, QRectF, Signal
from PySide6.QtGui import (QPainter, QColor, QLinearGradient,
                           QPainterPath, QTextCursor)
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QTextEdit,
    QPushButton, QComboBox, QLabel, QHBoxLayout, QFrame, QMessageBox,
    QFileDialog, QProgressBar, QGroupBox, QFormLayout, QLineEdit, QCheckBox, QSpinBox
)
from docx import Document

def format_transcription_text(raw_text: str) -> str:
    # 1. Fusionner les lignes
    text = raw_text.replace("\n", " ")
    # 2. Nettoyer ponctuation
    text = re.sub(r'\s*,\s*', ', ', text)
    text = re.sub(r'\s*\.\s*', '. ', text)
    text = re.sub(r'\s*\?\s*', '? ', text)
    text = re.sub(r'\s*!\s*', '! ', text)
    # 3. Découper en phrases
    sentences = re.split(r'(?<=[.!?])\s+', text)
    # 4. Reconstruction
    return "\n".join(s.strip() for s in sentences if s.strip())

class FileTranscribeThread(QThread):
    """Transcription d'un fichier en chunks, avec buffering de N phrases."""
    progress     = Signal(int, int)    # (current_chunk, total_chunks)
    segment      = Signal(str)         # paragraphes formatés
    audio_chunk  = Signal(object)      # pour l'affichage du waveform
    done         = Signal()

    def __init__(
        self,
        infile: str,
        model,
        model_name: str,
        lang: str,
        chunk_s: int = 30,
        spp: int = 3,
        beam_size: int = 5,
        best_of: int = 5
    ):
        super().__init__()
        self.infile    = infile
        self.model     = model
        self.model_name= model_name
        self.lang      = lang
        self.chunk_s   = chunk_s
        self.spp       = spp
        self.beam_size = beam_size
        self.best_of   = best_of

        self._abort    = False
        self.splitter  = re.compile(r'(?<=[\.\?\!])\s+')
        self.buffer    = []

    def run(self):
        # 1) Pré‐charger un modèle CPU de secours
        cpu_model = whisper.load_model(self.model_name, device="cpu")

        try:
            audio = whisper.load_audio(self.infile)
            sr, total = whisper.audio.SAMPLE_RATE, audio.shape[0]
            sz = self.chunk_s * sr
            chunks = math.ceil(total / sz)
            self.progress.emit(0, chunks)

            use_cpu = False  # flag pour basculer définitivement

            for i in range(chunks):
                if self._abort:
                    break

                start, end = i * sz, min((i + 1) * sz, total)
                chunk_data = audio[start:end]
                self.audio_chunk.emit(chunk_data)

                # 2) Choisir le modèle actif
                model = cpu_model if use_cpu else self.model

                try:
                    res = model.transcribe(
                        chunk_data,
                        language=self.lang,
                        beam_size=self.beam_size,
                        best_of=self.best_of,
                        fp16=False
                    )
                except RuntimeError as e:
                    msg = str(e).lower()
                    if "illegal memory access" in msg or "cuda" in msg:
                        # on passe en CPU pour la suite
                        use_cpu = True
                        # vider le cache sans risque de crash
                        try:
                            torch.cuda.empty_cache()
                        except Exception:
                            pass
                        # et relancer immédiatement en CPU
                        res = cpu_model.transcribe(
                            chunk_data,
                            language=self.lang,
                            beam_size=1,
                            best_of=1,
                            fp16=False
                        )
                    else:
                        # autre erreur -> on remonte
                        raise

                # 3) Bufferisation comme avant
                for seg in res["segments"]:
                    for ph in self.splitter.split(seg["text"].strip()):
                        if not ph:
                            continue
                        if len(self.buffer) == 0:
                            self.buffer.append(ph.capitalize())
                        else:
                            self.buffer.append(ph)
                        if len(self.buffer) >= self.spp:
                            para = " ".join(self.buffer)
                            self.segment.emit(para)
                            self.buffer.clear()

                self.progress.emit(i + 1, chunks)

            # flush final
            if self.buffer:
                self.segment.emit(" ".join(self.buffer))
                self.buffer.clear()

        except Exception:
            print("Erreur dans FileTranscribeThread :")
            traceback.print_exc()
        finally:
            self.done.emit()

    def stop(self):
        """Demande l’arrêt coopératif du thread."""
        self._abort = True


class WaveformWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.setMinimumHeight(100)
        self.waves = []
        self.target_waves = []
        self.animation_timer = QTimer()
        self.animation_timer.timeout.connect(self.update_waves)
        self.animation_timer.start(30)  # Faster updates for smoother animation
        self.is_recording = False
        self.transition_speed = 0.15  # Controls how fast waves transition

    def start_animation(self):
        self.is_recording = True
        self.waves = [0.1] * 30  # Start with small waves

    def stop_animation(self):
        self.is_recording = False
        self.target_waves = [0] * 30

    def update_audio_data(self, data):
        if len(data) > 0:
            normalized = np.abs(data) / np.max(np.abs(data) + 1e-10)
            if self.is_recording:
                # Create more varied wave heights with higher amplitude
                chunk_size = max(1, len(normalized)//30)
                self.target_waves = [
                    # Increased amplitude
                    normalized[i:i + chunk_size].mean() * 1.2
                    for i in range(0, len(normalized), chunk_size)
                ][:30]
                # Add more randomness for dynamic look (±30% variation)
                self.target_waves = [
                    w * (1 + np.random.uniform(-0.3, 0.3)) for w in self.target_waves]
            else:
                self.target_waves = [0] * 30
            self.update()

    def update_waves(self):
        if not self.waves:
            self.waves = [0] * 30
            self.target_waves = [0] * 30

        # More responsive wave transitions
        for i in range(len(self.waves)):
            if self.is_recording:
                # Enhanced dynamic variation
                target = self.target_waves[i] * \
                    (1 + np.sin(time.time() * 6 + i) * 0.15)
                target *= 1 + np.cos(time.time() * 4) * \
                    0.1  # Additional wave motion
            else:
                target = 0

            # Faster transitions
            self.waves[i] += (target - self.waves[i]) * 0.2

        self.update()

    def paintEvent(self, event):
        if not self.waves:
            return

        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)

        try:
            width = self.width()
            height = self.height()
            center_y = height / 2
            bar_width = width / (len(self.waves) * 1.5)
            max_height = height * 0.85  # Slightly higher bars

            # Green theme gradient
            gradient = QLinearGradient(0, 0, 0, height)
            if self.is_recording:
                # Vibrant green colors during recording
                gradient.setColorAt(0, QColor(46, 204, 113))  # Bright green
                gradient.setColorAt(0.5, QColor(39, 174, 96))  # Medium green
                gradient.setColorAt(1, QColor(33, 150, 83))  # Dark green
            else:
                # Subtle green when not recording
                gradient.setColorAt(0, QColor(46, 204, 113, 200))
                gradient.setColorAt(1, QColor(33, 150, 83, 200))

            painter.setPen(Qt.NoPen)
            painter.setBrush(gradient)

            for i, amplitude in enumerate(self.waves):
                x = width * i / len(self.waves)

                # Enhanced wave motion
                wave_effect = np.sin(time.time() * 4 + i * 0.5) * 0.08
                bar_height = max_height * (amplitude + wave_effect)

                if self.is_recording:
                    # Enhanced pulsing effect
                    pulse = 1 + np.sin(time.time() * 5) * 0.08
                    bar_height *= pulse

                rect = QRectF(
                    x + bar_width/2,
                    center_y - bar_height/2,
                    bar_width,
                    bar_height
                )

                # Green glow effect when recording
                if self.is_recording and amplitude > 0.1:
                    glow = QPainterPath()
                    glow.addRoundedRect(rect, bar_width/2, bar_width/2)
                    painter.fillPath(glow, QColor(46, 204, 113, 40))

                painter.drawRoundedRect(rect, bar_width/2, bar_width/2)

        finally:
            painter.end()

class ModelLoaderThread(QThread):
    loaded = Signal(object)   # émet le modèle une fois prêt
    error  = Signal(Exception)

    def __init__(self, model_name, device_str):
        super().__init__()
        self.model_name = model_name
        self.device_str = device_str

    def run(self):
        try:
            dev = torch.device("cuda" if self.device_str == "GPU" else "cpu")
            model = whisper.load_model(self.model_name, device=dev)
            self.loaded.emit(model)
        except Exception as e:
            self.error.emit(e)

class WhisperGUI(QMainWindow):
    update_text = Signal(str)
    add_newline = Signal()

    def __init__(self):
        super().__init__()
        self.current_transcription = ""  # Current transcription text
        self.history_text = []  # Array to store history
        self.current_segment_start = None  # Track start time of current segment
        self.loaded_file_path = None
        self.transcribing_file = False
        
        # Fichiers ouverts pour écriture temps réel
        self.txt_file = None
        self.docx_doc = None
        
        # Show startup message
        self.statusBar().showMessage("Application is starting...")

        self.progress_bar = QProgressBar()
        self.progress_bar.setMaximum(0)     # 0,0 pour mode indéterminé
        self.progress_bar.setVisible(False) # cachée par défaut
        self.statusBar().addPermanentWidget(self.progress_bar)
        
        self.init_ui()
        self.init_whisper()
        self.last_buffer_reset = time.time()
        self.update_text.connect(self.update_display)
        self.add_newline.connect(self._add_newline)

    def init_ui(self):
        # Layout chính
        main_layout = QVBoxLayout()

        # Frame cho controls
        controls_frame = QFrame()
        controls_frame.setFrameStyle(QFrame.Panel | QFrame.Raised)
        controls_layout = QHBoxLayout()

        # Model selection
        model_label = QLabel("Model:")
        self.model_combo = QComboBox()
        self.model_combo.addItems(["tiny", "base", "small", "medium", "large"])
        self.model_combo.currentTextChanged.connect(self.load_model)

        device_label = QLabel("Device:")
        self.device_combo = QComboBox()
        # si CUDA dispo, propose GPU, sinon juste CPU
        devices = (["GPU"] if torch.cuda.is_available() else []) + ["CPU"]
        self.device_combo.addItems(devices)
        if torch.cuda.is_available():
            # force la sélection de GPU par défaut
            self.device_combo.setCurrentText("GPU")
        controls_layout.addWidget(device_label)
        controls_layout.addWidget(self.device_combo)

        # Buttons
        self.record_button = QPushButton("Démarrer l'enregistrement")
        self.record_button.clicked.connect(self.toggle_recording)
        
        self.open_file_button = QPushButton("Ouvrir un fichier audio")
        self.open_file_button.clicked.connect(self.open_audio_file)

        # Add controls to layout
        controls_layout.addWidget(model_label)
        controls_layout.addWidget(self.model_combo)
        controls_layout.addStretch()
        controls_layout.addWidget(self.open_file_button)
        controls_layout.addWidget(self.record_button)
        controls_frame.setLayout(controls_layout)

        # Text display
        self.text_display = QTextEdit()
        self.text_display.setReadOnly(True)
        # Set dark theme
        self.text_display.setStyleSheet(
            "QTextEdit { background-color: #2b2b2b; color: white; }")

        # Waveform
        self.waveform = WaveformWidget()

        # Add everything to main layout
        main_layout.addWidget(controls_frame)
        main_layout.addWidget(self.waveform)
        main_layout.addWidget(self.text_display)

        # — NOUVEAU : zone "Enregistrement" —
        grp_export = QGroupBox("Options d'export en temps réel")
        form_export = QFormLayout()

        # — Mode expert et paramètres avancés —
        # Checkbox pour activer le groupe Expert
        self.chk_expert = QCheckBox("Mode expert")
        self.chk_expert.toggled.connect(self._on_expert_toggled)
        main_layout.addWidget(self.chk_expert)

        # Groupe Expert (caché par défaut)
        self.grp_exp = QGroupBox("Paramètres avancés")
        self.grp_exp.setVisible(False)
        form_exp = QFormLayout()

        # SpinBoxes pour chunk, spp, beam_size, best_of
        self.spn_chunk = QSpinBox()
        self.spn_chunk.setRange(5, 300)
        self.spn_chunk.setValue(30)
        form_exp.addRow("Chunk (s) :", self.spn_chunk)

        self.spn_spp = QSpinBox()
        self.spn_spp.setRange(1, 10)
        self.spn_spp.setValue(3)
        form_exp.addRow("Phrases/para :", self.spn_spp)

        self.spn_beam = QSpinBox()
        self.spn_beam.setRange(1, 10)
        self.spn_beam.setValue(5)
        form_exp.addRow("Beam size :", self.spn_beam)

        self.spn_best = QSpinBox()
        self.spn_best.setRange(1, 10)
        self.spn_best.setValue(5)
        form_exp.addRow("Best of :", self.spn_best)

        self.grp_exp.setLayout(form_exp)
        main_layout.addWidget(self.grp_exp)

        # TXT
        self.chk_save_txt  = QCheckBox("Enregistrer en TXT (temps réel)")
        self.chk_save_txt.setChecked(False)
        self.chk_save_txt.toggled.connect(self.toggle_txt_realtime)
        self.le_txt_path   = QLineEdit("transcription.txt")
        self.btn_txt_browse     = QPushButton("…")
        self.btn_txt_browse.setFixedWidth(30)
        self.btn_txt_browse.clicked.connect(
            lambda: self._browse(self.le_txt_path, save=True, filt="*.txt")
        )
        form_export.addRow(self.chk_save_txt,
                           self._hbox(self.le_txt_path, self.btn_txt_browse))

        # DOCX
        self.chk_save_docx = QCheckBox("Enregistrer en DOCX (temps réel)")
        self.chk_save_docx.setChecked(False)
        self.chk_save_docx.toggled.connect(self.toggle_docx_realtime)
        self.le_docx_path  = QLineEdit("transcription.docx")
        self.btn_docx_browse    = QPushButton("…")
        self.btn_docx_browse.setFixedWidth(30)
        self.btn_docx_browse.clicked.connect(
            lambda: self._browse(self.le_docx_path, save=True, filt="*.docx")
        )
        form_export.addRow(self.chk_save_docx,
                           self._hbox(self.le_docx_path, self.btn_docx_browse))

        grp_export.setLayout(form_export)
        main_layout.addWidget(grp_export)

        # Bouton d'enregistrement manuel (optionnel)
        self.btn_save = QPushButton("Enregistrer transcription manuellement")
        self.btn_save.clicked.connect(self.save_transcript_manual)
        main_layout.addWidget(self.btn_save, alignment=Qt.AlignRight)

        # Central widget
        central_widget = QWidget()
        central_widget.setLayout(main_layout)
        self.setCentralWidget(central_widget)

        # Window properties
        self.setWindowTitle("Whisper GUI")
        self.setGeometry(100, 100, 800, 600)

    @Slot(bool)
    def _on_expert_toggled(self, checked: bool):
        self.grp_exp.setVisible(checked)
        # Si vous voulez que la fenêtre redimensionne automatiquement :
        QTimer.singleShot(0, self.adjustSize)

    def toggle_txt_realtime(self, checked):
        """Active/désactive l'écriture temps réel en TXT"""
        if checked:
            try:
                self.txt_file = open(self.le_txt_path.text(), "w", encoding="utf-8")
                self.statusBar().showMessage("Écriture temps réel TXT activée", 2000)
            except Exception as e:
                QMessageBox.critical(self, "Erreur TXT", f"Impossible d'ouvrir le fichier: {e}")
                self.chk_save_txt.setChecked(False)
        else:
            if hasattr(self, 'txt_file') and self.txt_file:
                self.txt_file.close()
                self.txt_file = None
                self.statusBar().showMessage("Écriture temps réel TXT désactivée", 2000)

    def toggle_docx_realtime(self, checked):
        """Active/désactive l'écriture temps réel en DOCX"""
        if checked:
            try:
                self.docx_doc = Document()
                self.statusBar().showMessage("Écriture temps réel DOCX activée", 2000)
            except Exception as e:
                QMessageBox.critical(self, "Erreur DOCX", f"Erreur d'initialisation: {e}")
                self.chk_save_docx.setChecked(False)
        else:
            if hasattr(self, 'docx_doc') and self.docx_doc:
                try:
                    self.docx_doc.save(self.le_docx_path.text())
                    self.statusBar().showMessage("Document DOCX sauvegardé", 2000)
                except Exception as e:
                    QMessageBox.warning(self, "Avertissement", f"Erreur de sauvegarde DOCX: {e}")
                self.docx_doc = None

    def write_realtime(self, text):
        """Écrit le texte en temps réel dans les fichiers activés"""
        if self.txt_file:
            try:
                self.txt_file.write(text + "\n\n")
                self.txt_file.flush()  # Force l'écriture immédiate
                os.fsync(self.txt_file.fileno())
            except Exception as e:
                print(f"Erreur écriture TXT: {e}")

        if self.docx_doc:
            try:
                self.docx_doc.add_paragraph(text)
                # Sauvegarde périodique du DOCX
                self.docx_doc.save(self.le_docx_path.text())
            except Exception as e:
                print(f"Erreur écriture DOCX: {e}")

    def _hbox(self, widget, button):
        hb = QHBoxLayout()
        hb.addWidget(widget)
        hb.addWidget(button)
        return hb

    def _browse(self, line_edit: QLineEdit, save: bool, filt: str):
        """ Ouvre un QFileDialog getSave ou getOpen """
        dlg = (QFileDialog.getSaveFileName if save else QFileDialog.getOpenFileName)
        fn, _ = dlg(self, "Sélectionner un fichier", line_edit.text(), filt)
        if fn:
            line_edit.setText(fn)

    def save_transcript_manual(self):
        """Enregistrement manuel de la transcription (ancienne méthode)"""
        text = self.text_display.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "Rien à enregistrer",
                                "La zone de transcription est vide.")
            return

        # Découpage en paragraphes
        paras = [p.strip() for p in text.split("\n\n") if p.strip()]

        saved_files = []

        # TXT
        if not self.chk_save_txt.isChecked():  # Seulement si pas en temps réel
            try:
                with open(self.le_txt_path.text(), "w", encoding="utf-8") as f:
                    for p in paras:
                        f.write(p + "\n\n")
                saved_files.append("TXT")
            except Exception as e:
                QMessageBox.critical(self, "Erreur TXT", str(e))
                return

        # DOCX
        if not self.chk_save_docx.isChecked():  # Seulement si pas en temps réel
            try:
                doc = Document()
                for p in paras:
                    doc.add_paragraph(p)
                doc.save(self.le_docx_path.text())
                saved_files.append("DOCX")
            except Exception as e:
                QMessageBox.critical(self, "Erreur DOCX", str(e))
                return

        if saved_files:
            QMessageBox.information(self, "Terminé",
                                    f"Transcription enregistrée: {', '.join(saved_files)}")
        else:
            QMessageBox.information(self, "Info",
                                    "Les fichiers sont déjà en cours d'écriture temps réel")

    def init_whisper(self):
        self.recording = False
        self.audio_queue = queue.Queue()
        self.sample_rate = 16000
        self.channels = 1
        self.blocksize = int(self.sample_rate * 0.3)  # 0.3 giây mỗi chunk
        self.model = None
        self.process_thread = None
        self.stable_tokens = None
        self.unstable_tokens = None
        self.eos_token = None
        self.device = torch.device(
            'mps' if torch.backends.mps.is_available() else 'cpu')
        # Load model ngay khi khởi tạo
        self.load_model()

    def open_audio_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Sélectionner un fichier audio", "", "Audio (*.mp3 *.wav *.m4a *.ogg)"
        )
        if file_path:
            self.loaded_file_path = file_path
            self.statusBar().showMessage(f"Fichier chargé : {file_path}", 5000)

    @Slot(int, int)
    def _on_file_progress(self, current, total):
        """Met à jour la barre de progression"""
        if total > 0:
            progress = int((current * 100) / total)
            self.progress_bar.setValue(progress)
            self.statusBar().showMessage(f"Transcription en cours... {progress}% ({current}/{total})")

    @Slot(str)
    def _on_file_segment(self, text: str):
        """Reçoit chaque segment transcrit d'un fichier."""
        # Affichage dans l'interface
        self.text_display.append(text)
        
        # Écriture temps réel si activée
        self.write_realtime(text)

    def load_model(self):
        model_name = self.model_combo.currentText()
        device_str = self.device_combo.currentText()
        self.statusBar().showMessage(f"Téléchargement du modèle {model_name} sur {device_str}…")

        # Affiche la barre et désactive le combo
        self.progress_bar.setVisible(True)
        self.progress_bar.setMaximum(0)  # Mode indéterminé
        self.model_combo.setEnabled(False)
        self.device_combo.setEnabled(False)

        # Lance le thread
        self.loader = ModelLoaderThread(model_name, device_str)
        self.loader.loaded.connect(self.on_model_loaded)
        self.loader.error.connect(self.on_model_error)
        self.loader.start()

    @Slot(object)
    def on_model_loaded(self, model):
        self.model = model
        self.current_model_name = self.model_combo.currentText()
        self.statusBar().showMessage("Modèle chargé !", 3000)
        self.progress_bar.setVisible(False)
        self.model_combo.setEnabled(True)
        self.device_combo.setEnabled(True)

        # Réinitialisation éventuelle des tokens
        self.stable_tokens = None
        self.unstable_tokens = None
        self.eos_token = None

    @Slot(Exception)
    def on_model_error(self, err):
        QMessageBox.critical(self, "Erreur", f"Échec du téléchargement : {err}")
        self.progress_bar.setVisible(False)
        self.model_combo.setEnabled(True)
        self.device_combo.setEnabled(True)
    
    def process_audio(self):
        """Process audio data from the queue"""
        if self.model is None:
            print("Model not loaded. Please load the model first.")
            return

        audio_buffer = np.array([], dtype=np.float32)
        try:
            while self.recording:
                # Get audio data from queue
                if self.audio_queue.empty():
                    time.sleep(0.1)
                    continue

                # Lấy audio data mới và thêm vào buffer
                audio_data = self.audio_queue.get()
                audio_data = audio_data.flatten().astype(np.float32)
                audio_buffer = np.concatenate([audio_buffer, audio_data])

                # Giới hạn độ dài buffer để tránh quá tải
                max_buffer_size = self.sample_rate * 15  # 15 giây
                current_time = time.time()
                buffer_reset_time = 15
                # Reset buffer sau mỗi 15 giây
                if current_time - self.last_buffer_reset > buffer_reset_time:
                    # Save current transcription to history before reset
                    if self.current_transcription.strip():
                        # Format timestamps
                        end_time = datetime.fromtimestamp(current_time)
                        start_time = datetime.fromtimestamp(
                            self.current_segment_start or (current_time - buffer_reset_time))
                        timestamp = f"[{start_time.strftime('%H:%M:%S')}-{end_time.strftime('%H:%M:%S')}]"

                        # Add to history with timestamp
                        final_text = f"{timestamp} {self.current_transcription.strip()}"
                        self.history_text.append(final_text)
                        
                        # Écriture temps réel
                        self.write_realtime(final_text)

                    self.current_transcription = ""
                    self.current_segment_start = current_time  # Set start time for new segment
                    self.last_buffer_reset = current_time
                    self.add_newline.emit()  # Emit signal instead of direct modification
                    audio_buffer = audio_data  # Reset buffer
                elif len(audio_buffer) > max_buffer_size:
                    audio_buffer = audio_buffer[-max_buffer_size:]

                try:
                    # Transcription simplifiée avec OpenAI Whisper officiel
                    result = self.model.transcribe(
                        audio_buffer,
                        language="fr",
                        fp16=False
                    )

                    transcription = result["text"]

                    # Update the display with new text
                    self.current_transcription = transcription
                    self.update_text.emit(transcription)

                except Exception as e:
                    print(f"Error in transcription: {str(e)}")
                    traceback.print_exc()
                    continue

        except Exception as e:
            print(f"Error in process_audio: {str(e)}")
            traceback.print_exc()

    def toggle_recording(self):
        # 1) si on transcrit un fichier → on demande l’arrêt coopératif et on restaure immédiatement l’UI
        if self.transcribing_file:
            self.transcribing_file = False
            # on déconnecte pour éviter qu’à la fin du thread on ne ré-invoque on_file_done
            try:
                self.trans_file_thread.done.disconnect(self.on_file_done)
            except TypeError:
                pass

            self.trans_file_thread.stop()      # set _abort = True
            self.on_file_done()                # restaure tout de suite l’UI
            return

        # 2) sinon, si on est déjà en live mic → on arrête l’enregistrement
        if self.recording:
            self.stop_recording()
            return

        # 3) sinon, on démarre selon qu’on a un fichier chargé ou pas
        if self.loaded_file_path:
            self.start_file_transcription()
        else:
            self.start_recording()

    @Slot()
    def on_file_done(self):
        # Réactive boutons et exports
        self.transcribing_file = False
        # Réactiver exports TXT/DOCX
        self.chk_save_txt.setEnabled(True)
        self.chk_save_docx.setEnabled(True)
        self.btn_txt_browse.setEnabled(True)
        self.btn_docx_browse.setEnabled(True)

        self.chk_expert.setEnabled(True)
        self.grp_exp    .setEnabled(self.chk_expert.isChecked())

        self.open_file_button.setEnabled(True)
        self.record_button.setText("Démarrer l'enregistrement")
        # Cache la barre & réinitialise
        self.progress_bar.setVisible(False)
        self.loaded_file_path = None
        # Restauration de l’UI bloquée
        self.model_combo.setEnabled(True)
        self.device_combo.setEnabled(True)
        QMessageBox.information(self, "Terminé", "Transcription de fichier achevée.")
        self.waveform.stop_animation()
        try: self.trans_file_thread.audio_chunk.disconnect(self.waveform.update_audio_data)
        except: pass

    def start_file_transcription(self):
        # Réinitialisation de l’interface
        self.text_display.clear()
        self.history_text.clear()
        self.current_transcription = ""
        self.transcribing_file = True
        self.waveform.start_animation()

        # Désactiver le mode expert pendant la transcription
        self.chk_expert.setEnabled(False)
        self.grp_exp   .setEnabled(False)

        # Bloquer exports TXT/DOCX
        self.chk_save_txt  .setEnabled(False)
        self.chk_save_docx .setEnabled(False)
        self.btn_txt_browse .setEnabled(False)
        self.btn_docx_browse.setEnabled(False)

        # Préparer la barre de progression
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.progress_bar.setMaximum(100)

        # Désactiver les contrôles durant la transcription
        self.open_file_button.setEnabled(False)
        self.record_button.setText("Arrêter transcription fichier")
        self.model_combo.setEnabled(False)
        self.device_combo.setEnabled(False)

        # --- Calcul dynamique de la durée de chunk ---
        # On charge l’audio pour en mesurer la durée totale
        audio = whisper.load_audio(self.loaded_file_path)
        sr = whisper.audio.SAMPLE_RATE
        total_seconds = audio.shape[0] / sr

        if total_seconds > 2 * 3600:
            chunk_duration = 120
        elif total_seconds > 3600:
            chunk_duration = 60
        elif total_seconds > 30 * 60:
            chunk_duration = 45
        else:
            chunk_duration = self.spn_chunk.value()

        # --- Lancement du thread de transcription fichier ---
        self.trans_file_thread = FileTranscribeThread(
            infile    = self.loaded_file_path,
            model     = self.model,
            model_name= self.current_model_name,
            lang      = "fr",
            chunk_s   = chunk_duration,
            spp       = self.spn_spp.value(),
            beam_size = self.spn_beam.value(),
            best_of   = self.spn_best.value()
        )
        self.trans_file_thread.audio_chunk.connect(self.waveform.update_audio_data)
        self.trans_file_thread.progress   .connect(self._on_file_progress)
        self.trans_file_thread.segment    .connect(self._on_file_segment)
        self.trans_file_thread.done       .connect(self.on_file_done)
        self.trans_file_thread.start()

    def start_recording(self):
        if not self.model or self.model_combo.isEnabled()==False:
            QMessageBox.warning(self, "Patientez", "Le modèle n'est pas encore chargé.")
            return

        self.text_display.clear()
        self.history_text.clear()
        self.current_transcription = ""
        # désactive exports
        self.chk_save_txt.setEnabled(False)
        self.chk_save_docx.setEnabled(False)

        # Désactiver le Mode expert pendant la transcription de fichier
        self.chk_expert.setEnabled(False)
        self.grp_exp    .setEnabled(False)

        self.recording = True
        self.record_button.setText("Arrêter l'enregistrement")
        self.waveform.start_animation()

        # Start processing thread
        self.process_thread = threading.Thread(target=self.process_audio)
        self.process_thread.start()

        self.model_combo.setEnabled(False)
        self.device_combo.setEnabled(False)

        # Start audio input stream
        self.stream = sd.InputStream(
            samplerate=self.sample_rate,
            channels=self.channels,
            callback=self.audio_callback,
            blocksize=self.blocksize
        )
        self.stream.start()

    def stop_recording(self):
        self.recording = False
        # réactive exports
        self.chk_save_txt.setEnabled(True)
        self.chk_save_docx.setEnabled(True)
        self.record_button.setText("Démarrer l'enregistrement")
        self.waveform.stop_animation()

        # … après avoir réactivé exports et controls
        self.chk_expert.setEnabled(True)
        self.grp_exp    .setEnabled(self.chk_expert.isChecked())


        self.model_combo.setEnabled(True)
        self.device_combo.setEnabled(True)

        if hasattr(self, 'stream'):
            self.stream.stop()
            self.stream.close()

        if self.process_thread:
            self.process_thread.join()

        # Sauvegarde finale du segment en cours
        if self.current_transcription.strip():
            current_time = time.time()
            end_time = datetime.fromtimestamp(current_time)
            start_time = datetime.fromtimestamp(
                self.current_segment_start or (current_time - 15))
            timestamp = f"[{start_time.strftime('%H:%M:%S')}-{end_time.strftime('%H:%M:%S')}]"
            
            final_text = f"{timestamp} {self.current_transcription.strip()}"
            self.history_text.append(final_text)
            
            # Écriture temps réel
            self.write_realtime(final_text)

    def audio_callback(self, indata, frames, time, status):
        """Callback for audio input"""
        if status:
            print(status)
        # Add the new audio data to the queue
        self.audio_queue.put(indata.copy())
        self.waveform.update_audio_data(indata.copy())

    def merge_text(self, text1, text2):
        """
        Merge two texts intelligently, keeping the context and avoiding duplicates
        """
        if not text1:
            return text2

        # Tìm phần chung dài nhất giữa cuối text1 và đầu text2
        words1 = text1.lower().split()
        words2 = text2.lower().split()

        max_overlap = 0
        overlap_pos = 0

        # Tìm vị trí overlap tốt nhất
        for i in range(len(words1)):
            for j in range(len(words2)):
                k = 0
                while (i + k < len(words1) and
                       k < len(words2) and
                       words1[i + k] == words2[k]):
                    k += 1
                if k > max_overlap:
                    max_overlap = k
                    overlap_pos = i

        if max_overlap > 0:
            # Kết hợp text với phần overlap
            result = " ".join(words1[:overlap_pos] + words2)
        else:
            # Nếu không có overlap, nối trực tiếp
            result = text1 + " " + text2

        return result

    def _add_newline(self):
        if len(self.text_display.toPlainText().strip()) > 0:
            self.text_display.append("")

    def update_display(self, text):
        """
        Récupère l'historique + le segment courant, formate l'ensemble
        et met à jour le QTextEdit en temps réel.
        """
        # 1) Concatène l’historique et le nouveau segment
        full = ""
        if self.history_text:
            full = "\n".join(self.history_text) + "\n\n"
        full += text

        # 2) Formate tout de suite
        formatted = format_transcription_text(full)

        # 3) Affiche
        self.text_display.setPlainText(formatted)

        # 4) Replace le curseur à la fin
        cursor = self.text_display.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.text_display.setTextCursor(cursor)

    def closeEvent(self, event):
        # Arrête proprement l’enregistrement live
        if getattr(self, 'recording', False):
            self.stop_recording()

        # Si on est en cours de transcription de fichier, on arrête et on attend le thread
        if getattr(self, 'transcribing_file', False) and hasattr(self, 'trans_file_thread'):
            self.trans_file_thread.stop()
            self.trans_file_thread.wait()

        # Si le loader de modèle tourne toujours, on l’arrête aussi
        if hasattr(self, 'loader') and self.loader.isRunning():
            self.loader.terminate()
            self.loader.wait()

        # Ferme les fichiers temps-réel s’ils sont ouverts
        if getattr(self, 'txt_file', None):
            self.txt_file.close()
        if getattr(self, 'docx_doc', None):
            try:
                self.docx_doc.save(self.le_docx_path.text())
            except:
                pass

        super().closeEvent(event)



def main():
    app = QApplication(sys.argv)
    window = WhisperGUI()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()